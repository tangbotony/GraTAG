import datetime
import os
import shutil
import uuid
from http import HTTPStatus
from datetime import timedelta

import pypandoc
from controllers import (ALLOWED_EXTENSIONS, ALLOWED_IMAGE, IMAGE_URL,
                         UPLOAD_FILE_DIR)
from controllers.utils import (change_to_meta, get_path, validate_image,
                               validator_parent_id)
from flask import Blueprint, jsonify, make_response
from flask_jwt_extended import get_jwt_identity, jwt_required
from models.document import FILE_TYPE_DOCUMENT, Document
from models.image import Image
from models.upload_file import UploadDocument
from docx import Document as Document_docx
from pdf2docx import Converter
from webargs import fields
from webargs.flaskparser import use_args
from werkzeug.utils import secure_filename

router_document = Blueprint('document', __name__)

# 新建文档
@router_document.route('/api/document', methods=['POST'])
@jwt_required()
@use_args({"text": fields.Str(required=True), "name": fields.Str(required=True), "plain_text": fields.Str(),
           "parent_id": fields.Str(required=True), "reference": fields.List(cls_or_instance=fields.Dict(), required=False)},location="json")
def create_document(data):
    user_id = get_jwt_identity()
    text, name, parent_id, plain_text, reference = data.get('text'), data.get('name'), data.get("parent_id"), data.get("plain_text"), data.get("reference")
    if not validator_parent_id(parent_id, user_id):
        return jsonify({'message': "not found parent id in database"}), HTTPStatus.BAD_REQUEST

    now = (datetime.datetime.now() + timedelta(hours=8)).strftime("%Y-%m-%d %H:%M:%S")
    id = str(uuid.uuid4())
    path = get_path(parent_id, id, name, user_id)

    doc = Document(_id=id, user_id=user_id, name=name, text=text, plain_text=plain_text, parent_id=parent_id, type=FILE_TYPE_DOCUMENT,
                    path=path, create_time=now, update_time=now, extend_data={}, reference=reference).save()  
    if not doc:
        return jsonify({'message': "create doc in mongodb failed"}), HTTPStatus.INTERNAL_SERVER_ERROR
    return jsonify({'message': "", 'doc_id': doc._id}), HTTPStatus.OK
  

# 修改保存文档
@router_document.route('/api/document', methods=['PUT'])
@use_args({"id": fields.Str(required=True), "name": fields.Str(), "plain_text": fields.Str(),
           "parent_id": fields.Str(), "text": fields.Str(), "reference": fields.List(cls_or_instance=fields.Dict(), required=False), 
           "is_quote": fields.Bool()}, location="json")
@jwt_required()
def update_document(data):
    user_id = get_jwt_identity()
    now = (datetime.datetime.now() + timedelta(hours=8)).strftime("%Y-%m-%d %H:%M:%S")
    doc = Document.objects(_id=data.get("id"), user_id=user_id, is_trash=False, is_delete=False,
                           ).first()
    if not doc:
        return jsonify({'message': "未在您的文件库中找到该文件 或 您无该篇文档的操作权限"}), HTTPStatus.BAD_REQUEST
    
    # parent id 校验
    if "parent_id" in data:
        if not validator_parent_id(data.get("parent_id"), user_id):
            return jsonify({'message': "not found parent id in database"}), HTTPStatus.BAD_REQUEST
        path = get_path(data.get("parent_id"), data.get("id"), doc.name, user_id)
        data['path'] = path

    data['update_time'] = now
    data.pop("id")  
    status = doc.update(**data)
    return jsonify({'message': "", 'status': status}), HTTPStatus.OK


# 查阅文档
@router_document.route('/api/document/<id>', methods=['GET'])
@jwt_required()
def get_document(id):
    user_id = get_jwt_identity()

    doc = Document.objects(_id=id, user_id=user_id, is_trash=False, is_delete=False).first()
    if not doc:
        return jsonify({'message': "未在您的文件库中找到该文件 或 您无该篇文档的操作权限"}), HTTPStatus.INTERNAL_SERVER_ERROR
    return jsonify({'message': "", 'doc': doc}), HTTPStatus.OK


# 文档放入废纸篓
@router_document.route('/api/document/<id>', methods=['DELETE'])
@jwt_required()
def trash_document(id):
    user_id = get_jwt_identity()
    trash_time = datetime.datetime.now() + datetime.timedelta(days=30)
    trash_time = trash_time.strftime("%Y-%m-%d %H:%M:%S")

    doc = Document.objects(_id=id, user_id=user_id, is_trash=False, is_delete=False).first()
    if not doc:
        return jsonify({'message': "未在您的文件库中找到该文件 或 您无该篇文档的操作权限"}), HTTPStatus.INTERNAL_SERVER_ERROR
    status = doc.update(is_trash=True, trash_time=trash_time, trash_root=True)
    return jsonify({'message': "", 'status': status}), HTTPStatus.OK


# 文档搜索
@router_document.route('/api/search/document', methods=['POST'])
@jwt_required()
@use_args({"search_text": fields.Str(required=True), "parent_id": fields.Str()}, location="json")
def search_document(data):
    user_id = get_jwt_identity()

    search_text, parent_id = data.get("search_text"), data.get("parent_id", "-1")
    docs = Document.objects(is_trash=False, user_id=user_id).all()

    # 字符匹配搜索
    res = []
    for item in docs:
        if search_text not in item.name:
            continue
        for (id, name) in item.path[:-1]:
            if id == parent_id:
                res.append(item)
                break
    meta_list = change_to_meta(res)
    meta_list.sort()
    meta_list = [item.__dict__ for item in meta_list]
    return jsonify({'message': "", 'meta': meta_list}), HTTPStatus.OK


# 文档上传，只支持txt，docx，pdf
@router_document.route('/api/upload', methods=['POST'])
@jwt_required()
@use_args({"file": fields.Field(required=True)}, location="files")
def upload_document(data):
    user_id = get_jwt_identity()
    file = data['file']
    if not allowed_file(file.filename):
        return jsonify({'message': "not allowed file, only support pdf, docx, txt file"}), HTTPStatus.INTERNAL_SERVER_ERROR

    now = (datetime.datetime.now() + timedelta(hours=8)).strftime("%Y-%m-%d %H:%M:%S")
    id = str(uuid.uuid4())
    base_path = os.path.join(UPLOAD_FILE_DIR, "{}_{}".format(user_id, id))
    os.makedirs(base_path)
    
    # 原始文件暂存本地用于格式转换
    filename = secure_filename(file.filename)
    file.save(os.path.join(base_path, filename))
    suffix = ALLOWED_EXTENSIONS.get(filename.rsplit('.', 1)[1].lower())
    raw_filename = filename.rsplit('.', 1)[0].lower()

    # 格式转换，提取文本
    source_file, outputfile = os.path.join(base_path, filename), os.path.join(base_path, "{}.txt".format(raw_filename))
    change_format_to_txt(suffix, source_file, outputfile, base_path)
    
    # 原始文件和文本内容存储到mongodb
    text = ""
    with open(os.path.join(base_path, "{}.txt".format(raw_filename)), 'r') as f:
        text = f.read()
    doc = UploadDocument(_id=id, user_id=user_id, name=raw_filename, text=text, format=suffix, create_time=now)
    with open(os.path.join(base_path, filename), 'rb') as fd:
        doc.file.put(fd, content_type=suffix)
    doc = doc.save()

    # 删除本地文件
    shutil.rmtree(base_path)

    return jsonify({'message': "", 'doc_id': doc._id}), HTTPStatus.OK

def post_process(text):
    text = text.replace("\ufffd"," ")
    return text

# 文档解析，支持txt，docx, doc，pdf
@router_document.route('/api/analyse', methods=['POST'])
@jwt_required()
@use_args({"file": fields.Field(required=True)}, location="files")
def analyse_document(data):
    user_id = get_jwt_identity()
    file = data['file']
    if not allowed_file(file.filename):
        return jsonify({'message': "not allowed file, only support pdf, docx, doc, txt file"}), HTTPStatus.INTERNAL_SERVER_ERROR
    try:
        now = (datetime.datetime.now() + timedelta(hours=8)).strftime("%Y-%m-%d %H:%M:%S")
        id = str(uuid.uuid4())
        base_path = os.path.join(UPLOAD_FILE_DIR, "{}_{}".format(user_id, id))
        os.makedirs(base_path)
        
        # 原始文件暂存本地用于格式转换
        filename = secure_filename(file.filename)
        if '.' not in filename:
            filename = file.filename
        file.save(os.path.join(base_path, filename))
        suffix = ALLOWED_EXTENSIONS.get(filename.rsplit('.', 1)[1].lower())
        raw_filename = filename.rsplit('.', 1)[0].lower()
        if suffix == 'pdf':
            pdf_text = extract_text_from_pdf(os.path.join(base_path, filename))
            if pdf_text:
                return jsonify({'res': pdf_text, 'err': ""}), HTTPStatus.OK

        # 格式转换，提取文本
        source_file, outputfile = os.path.join(base_path, filename), os.path.join(base_path, "{}.txt".format(raw_filename))
        change_format_to_txt(suffix, source_file, outputfile, base_path)
        
        # 原始文件和文本内容存储到mongodb
        text = ""
        with open(os.path.join(base_path, "{}.txt".format(raw_filename)), 'r', encoding='utf-8') as f:
            text = f.read()
            text = post_process(text)
            
        # 删除本地文件
        shutil.rmtree(base_path)
    except Exception as e:
        return jsonify({'res': "", 'err': f"文档解析失败: {e}"}), HTTPStatus.INTERNAL_SERVER_ERROR

    return jsonify({'res': text, 'err': ""}), HTTPStatus.OK


from material import save_img_to_oss
from io import BytesIO
IMAGE_OSS_KEY='img'

# image上传
@router_document.route('/api/upload/image', methods=['POST'])
@jwt_required()
@use_args({"image": fields.Field(required=True)}, location="files")
def upload_image(data):

    file = data["image"]
    filename = file.filename
    if not validate_image(filename):
        return jsonify({'message': "not allowed image, only support {}".format(ALLOWED_IMAGE)}), HTTPStatus.BAD_REQUEST
    id = str(uuid.uuid4())
    suffix = filename.rsplit('.', 1)[1].lower()
    filename = "{}.{}".format(id, suffix)   
    file_content = BytesIO(file.stream.read())
    save_img_to_oss(filename, file_content, IMAGE_OSS_KEY)
    return jsonify({'message': "", 'image_url': "{}/{}/{}".format(IMAGE_URL, IMAGE_OSS_KEY, filename)}), HTTPStatus.OK



def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def change_format_to_txt(suffix, source_file, outputfile, base_path):
    if suffix == "docx":
        pypandoc.convert_file(source_file=source_file, to='plain', 
                      format=suffix, outputfile=outputfile, 
                      encoding="utf-8", extra_args=['--extract-media={}'.format(base_path)])
    if suffix == "doc":
        docx_file = os.path.join(base_path, "doc_to_docx.docx")
        Document_docx(source_file).save(docx_file)
        pypandoc.convert_file(source_file=docx_file, to='plain', 
                      format="docx", outputfile=outputfile, 
                      encoding="utf-8", extra_args=['--extract-media={}'.format(base_path)])
    elif suffix == "pdf":
        docx_file = os.path.join(base_path, "pdf_test.docx")
        cv = Converter(source_file)
        cv.convert(docx_file) 
        cv.close()
        pypandoc.convert_file(source_file=docx_file, to='plain', 
                      format="docx", outputfile=outputfile, 
                      encoding="utf-8", extra_args=['--extract-media={}'.format(base_path)])
    return



from io import StringIO
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from concurrent.futures import ProcessPoolExecutor, ThreadPoolExecutor
def miner(page_number, pdf_path):
    output_string = StringIO()
    with open(pdf_path, 'rb') as in_file:
        parser = PDFParser(in_file)
        doc = PDFDocument(parser)
        rsrcmgr = PDFResourceManager()
        device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        for pageNum, page in enumerate(PDFPage.create_pages(doc)):
            if pageNum == page_number:
                interpreter.process_page(page)
                break

    return output_string.getvalue()

def extract_text_from_pdf1(pdf_path):
    with open(pdf_path, 'rb') as in_file:
        try:
            doc = PDFDocument(PDFParser(in_file))

            with ThreadPoolExecutor(max_workers=10) as executor:
                futures = [executor.submit(miner, pageNum, pdf_path) for pageNum in range(len(list(PDFPage.create_pages(doc))))]
                results = [future.result() for future in futures]

            text = "".join(results).replace('\t', ' ')
        except Exception as e:
            print(f"Error processing the PDF file: {e}")
            return None
    return text

# def extract_text_from_pdf(pdf_path):
#     output_string = StringIO()
#     with open(pdf_path, 'rb') as in_file:
#         parser = PDFParser(in_file)
#         doc = PDFDocument(parser)
#         rsrcmgr = PDFResourceManager()
#         device = TextConverter(rsrcmgr, output_string, laparams=LAParams())
#         interpreter = PDFPageInterpreter(rsrcmgr, device)
#         for page in PDFPage.create_pages(doc):
#             interpreter.process_page(page)
#     return output_string.getvalue()

import pdfplumber
def extract_text_from_pdf3(pdf_path):
    with pdfplumber.open(pdf_path) as pdf:
        # 遍历每一页
        text = ''
        for page in pdf.pages:
            text += page.extract_text()
        
    return text

import pymupdf
def extract_text_from_pdf(pdf_path):
    with pymupdf.open(pdf_path) as pdf:
        # 遍历每一页
        text = ''
        for page in pdf:
            text += page.get_text()
        
    return text

import pypdfium2

